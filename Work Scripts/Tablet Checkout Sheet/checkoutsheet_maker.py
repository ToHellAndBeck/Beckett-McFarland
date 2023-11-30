import os
import cv2
from PIL import Image
import pytesseract
import re
import openpyxl
from docx import Document
from pyzbar.pyzbar import decode
from docx.shared import Pt, RGBColor

# Set the path to the Tesseract executable (update with your path if necessary)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
if __name__ == "__main__":
    def extract_numbers_from_image(image_path):
        try:
            # Open the image file
            image = Image.open(image_path)
            
            # Use Tesseract to do OCR on the image with adjusted parameters
            text = pytesseract.image_to_string(image, config='--psm 6 --oem 3')
            
            # Extract only numbers using regular expression
            numbers = re.findall(r'\d+', text)
            
            return numbers
        except Exception as e:
            return f"Error: {e}"

    def extract_numbers_from_qr_code(image_path):
        try:
            # Open the image file
            image = Image.open(image_path)

            # Use pyzbar to decode QR code
            decoded_objects = decode(image)
            
            if decoded_objects:
                # Extract the decoded information from the first QR code
                decoded_info = decoded_objects[0].data.decode('utf-8')
                return decoded_info
            else:
                return None

        except Exception as e:
            return f"Error: {e}"
    """
    def combine_first_three_numbers(numbers):
        # Take the first three numbers and concatenate them into a single long string
        return ''.join(numbers[:3])
    """
    def retrieve_data_from_excel(combined_string, excel_file_path, sheet_name, columns_to_extract):
        
        try:
            # Load the Excel workbook
            workbook = openpyxl.load_workbook(excel_file_path, data_only=True)
            
            # Select the sheet by name
            sheet = workbook[sheet_name]
            
            # Iterate through the rows and find the matching value in the first column
            for row in sheet.iter_rows(min_row=2, values_only=True):
                if row[0] and str(row[0]) == combined_string:
                    # If a match is found, retrieve data from specified columns
                    row_data = [row[col - 1] for col in columns_to_extract]  # Adjusted indexing
                    return row_data
            
            return None  # Return None if no match is found
        
        except Exception as e:
            return f"Error: {e}"

    def remove_word_from_data(data, word_to_remove):
        # Remove the specified word from each element in the data list
        return [value.replace(word_to_remove, '') if isinstance(value, str) else value for value in data]

    def replace_placeholders_in_docx(docx_path, placeholders):
        doc = Document(docx_path)
        
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for placeholder, replacement in placeholders.items():
                    if placeholder in run.text:
                        print("Info added to document!")
                        run.text = run.text.replace(placeholder, replacement)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for placeholder, replacement in placeholders.items():
                                if placeholder in run.text:
                                    print("Info added to document!")
                                    run.text = run.text.replace(placeholder, replacement)

        return doc

    def format_combined_string(combined_string):
        # Remove all non-numeric characters from the phone number
        clean_number = re.sub('[^0-9]+', '', combined_string)

        # Format the phone number using regex
        formatted_number = re.sub("(\\d{3})(\\d{3})(\\d+)", "(\\1) \\2-\\3", clean_number)

        return formatted_number


# Rest of the code remains unchanged
# ...

# Inside the main block

    # Replace 'your_qrcodes_folder' with the path to your folder containing QR code images
    qrcodes_folder = r"L:\Rollout\Coordinator\Tablet Checkout Sheet Scripts\QR Code Loader"
    
    # Iterate through QR code images in the folder
    for filename in os.listdir(qrcodes_folder):
        if filename.endswith(".png"):
            qr_code_path = os.path.join(qrcodes_folder, filename)
            print(f"Processing QR code: {qr_code_path}")

            # Replace 'your_combined_string' with the actual combined string
            extracted_numbers = extract_numbers_from_qr_code(qr_code_path)
            
            if extracted_numbers:
                combined_string = extracted_numbers

                formatted_number = format_combined_string(combined_string)
                print("TABLET PHONE NUMBER: " + formatted_number)

                if combined_string:
                    print("Searching the Excel Sheet for: " + combined_string)
                    
                    # Replace 'your_excel_file.xlsx', 'Sheet1', and [2, 3, 4] with your actual Excel file path, sheet name, and columns to extract
                    excel_file_path = r"L:\Rollout\Tablets\Master Tablet 11-26-2023.xlsx"
                    sheet_name = 'CURRENT TABLETS'
                    columns_to_extract = [5, 6]
                    
                    row_data = retrieve_data_from_excel(combined_string, excel_file_path, sheet_name, columns_to_extract)
                    
                    if row_data is not None:
                        # Replace 'word_to_remove' with the actual word you want to remove
                        word_to_remove = 'Samsung '
                        
                        # Remove the specified word from the data
                        row_data_after_removal = remove_word_from_data(row_data, word_to_remove)
                        print(f"IMEI: {row_data_after_removal[0]}")
                        print(f"MODEL: {row_data_after_removal[1]}")
                        
                        # Define placeholders and their replacements
                        placeholders = {
                            "[IMEI NUMBER]": str(row_data_after_removal[0]),
                            "[MODEL]": str(row_data_after_removal[1]),
                            "[PHONE NUMBER]": str(formatted_number)
                        }

                        # Replace placeholders in the Word document
                        template_path = r"L:\Rollout\Coordinator\Tablet Checkout Sheet Scripts\Template\checkout_template.docx"
                        modified_doc = replace_placeholders_in_docx(template_path, placeholders)
                        qrcodes_folder=r"L:\Rollout\Coordinator\Tablet Checkout Sheet Scripts\QR Code Loader"
                        # Save the modified document as a new file
                        output_doc_path = os.path.join(r"L:\Rollout\Coordinator\Tablet Checkout Sheet Scripts\Output", f"output_{filename.replace('.png', '.docx')}")
                        modified_doc.save(output_doc_path)
                        print(f"Document saved: {output_doc_path}")

                    else:
                        user_input = input("No match found in the Excel sheet.")

                else:
                    print("Number extraction failed.")