from PIL import Image
import pytesseract
import re
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor
import string
import os
import random

# Set the path to the Tesseract executable (update with your path if necessary)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def save_as_separate_document(doc, output_folder, prefix="output"):
    random_suffix = ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))
    output_path = os.path.join(output_folder, f"{prefix}_{random_suffix}.docx")
    doc.save(output_path)
    print(f"Document saved: {output_path}")

def extract_numbers_from_image(image_path):
    try:
        # Open the image file
        image = Image.open(image_path)
        
        # Use Tesseract to do OCR on the image
        text = pytesseract.image_to_string(image)
        
        # Extract only numbers using regular expression
        numbers = re.findall(r'\d+', text)
        
        return numbers
    except Exception as e:
        return f"Error: {e}"

def combine_first_three_numbers(numbers):
    # Take the first three numbers and concatenate them into a single long string
    return ''.join(filter(str.isdigit, numbers))

def retrieve_data_from_excel(combined_string, excel_text_file_folder, sheet_name, columns_to_extract):
    try:
        # Load the Excel workbook
        workbook = openpyxl.load_workbook(excel_text_file_folder)
        
        # Select the sheet by name
        sheet = workbook[sheet_name]
        
        # Iterate through the rows and find the matching value in the first column
        for row in sheet.iter_rows(min_row=2, values_only=True):
            if row[0] and str(row[0]) == combined_string:
                # If a match is found, retrieve data from specified columns
                row_data = [row[col - 1] for col in columns_to_extract]  # Adjusted indexing
                return row_data
        
        return None  # Return None if no match is found
    
    except Exception as e:
        return f"Error: {e}"

def remove_word_from_data(data, word_to_remove):
    # Remove the specified word from each element in the data list
    return [value.replace(word_to_remove, '') if isinstance(value, str) else value for value in data]

def replace_placeholders_in_docx(docx_path, placeholders):
    doc = Document(docx_path)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for placeholder, replacement in placeholders.items():
                if placeholder in run.text:
                    print("Info added to document!")
                    run.text = run.text.replace(placeholder, replacement)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, replacement in placeholders.items():
                            if placeholder in run.text:
                                print("Info added to document!")
                                run.text = run.text.replace(placeholder, replacement)

    doc.save(docx_path)  # Save the changes to the original file
    print("Document saved.")
def extract_numbers_from_text(text_file_folder):
    try:
        with open(text_file_folder, 'r') as file:
            return [line.strip() for line in file.readlines()]
    except Exception as e:
        return f"Error: {e}"
if __name__ == "__main__":
    # Replace 'your_text_file.txt' with the path to your text file
    text_file_folder = r"C:\Users\beckett.mcfarland\Downloads\image-to-text"
    output_folder = r"C:\Users\beckett.mcfarland\Desktop\Beckett-McFarland\Work Scripts\Tablet Checkout Sheet\Output"

    for filename in os.listdir(text_file_folder):
        if filename.lower().endswith('.txt'):
            text_text_file_folder = os.path.join(text_file_folder, filename)
            extracted_numbers = extract_numbers_from_text(text_text_file_folder)

            for extracted_number in extracted_numbers:
                if extracted_number:
                    def format_combined_string(combined_string):
                        # Remove all non-numeric characters from the phone number
                        clean_number = re.sub('[^0-9]+', '', combined_string)

                        # Format the phone number using regex
                        formatted_number = re.sub("(\\d{3})(\\d{3})(\\d+)", "(\\1) \\2-\\3", clean_number)

                        return formatted_number

                    combined_string = combine_first_three_numbers(extracted_number)

                    formatted_number = format_combined_string(combined_string)
                    print(f"TABLET PHONE NUMBER ({filename}): {formatted_number}")

                    if combined_string:
                        print(f"Searching the Excel Sheet for ({filename}): {combined_string}")

                        excel_text_file_folder = r"L:\Rollout\Tablets\Master Tablet 11-26-2023.xlsx"
                        sheet_name = 'CURRENT TABLETS'
                        columns_to_extract = [5, 6]

                        row_data = retrieve_data_from_excel(combined_string, excel_text_file_folder, sheet_name, columns_to_extract)

                        if row_data is not None:
                            template_path = r"C:\Users\beckett.mcfarland\Desktop\Beckett-McFarland\Work Scripts\Tablet Checkout Sheet\Template\checkout_template.docx"

                            # Load the template document
                            template_doc = Document(template_path)

                            # Update the first paragraph text
                            template_doc.paragraphs[0].text = f"TABLET PHONE NUMBER ({filename}): {formatted_number}"

                            # Update the first table cell text
                            template_doc.tables[0].cell(0, 0).paragraphs[0].text = f"IMEI ({filename}): {row_data[0]}"
                            template_doc.tables[0].cell(0, 1).paragraphs[0].text = f"MODEL ({filename}): {row_data[1]}"

                            placeholders = {
                                "[IMEI NUMBER]": str(row_data[0]),
                                "[MODEL]": str(row_data[1]),
                                "[PHONE NUMBER]": str(formatted_number)
                            }

                            replace_placeholders_in_docx(template_doc, placeholders)

                            # Save the modified template as a new document
                            save_as_separate_document(template_doc, output_folder, prefix=f"output_{filename[:-4]}")
                        else:
                            print(f"No match found in the Excel sheet for ({filename}).")
                    else:
                        print(f"Combining numbers failed for ({filename}).")
                else:
                    print(f"Failed to extract number from text file: {filename}")