import os
from PIL import Image
from docx import Document
from pyzbar.pyzbar import decode

def extract_info_from_qr_code(image_path):
    try:
        image = Image.open(image_path)
        decoded_objects = decode(image)
        
        if decoded_objects:
            decoded_info = decoded_objects[0].data.decode('utf-8')
            # Splitting the decoded info into lines
            lines = decoded_info.split("\n")
            info_dict = {}
            for line in lines:
                if ": " in line:
                    key, value = line.split(": ", 1)  # Splitting each line into key-value pair
                    info_dict[key] = value
            return info_dict
        else:
            return None

    except Exception as e:
        return f"Error: {e}"

def replace_placeholders_in_docx(docx_path, placeholders):
    doc = Document(docx_path)
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for placeholder, replacement in placeholders.items():
                if placeholder in run.text:
                    print("Info added to document!")
                    run.text = run.text.replace(placeholder, replacement)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, replacement in placeholders.items():
                            if placeholder in run.text:
                                print("Info added to document!")
                                run.text = run.text.replace(placeholder, replacement)

    return doc

if __name__ == "__main__":
    qrcodes_folder =r"C:\Users\beckett.mcfarland\Desktop\Beckett-McFarland\Work Scripts\Tablet Checkout Sheet\QR Codes"
    template_path = r"C:\Users\beckett.mcfarland\Desktop\Beckett-McFarland\Work Scripts\Tablet Checkout Sheet\Template\checkout_template.docx"
    output_folder_path = r"C:\Users\beckett.mcfarland\Desktop\Beckett-McFarland\Work Scripts\Tablet Checkout Sheet\Output"

    for filename in os.listdir(qrcodes_folder):
        if filename.endswith(".png"):
            qr_code_path = os.path.join(qrcodes_folder, filename)
            print(f"Processing QR code: {qr_code_path}")

            extracted_info = extract_info_from_qr_code(qr_code_path)
            
            if isinstance(extracted_info, dict):  # Check if extracted_info is a dictionary
                # Extract phone number from file name (assuming it's before the extension)
                phone_number = filename.split(".")[0]

                # Prepare placeholders with extracted and filename data
                placeholders = {
                    "[IMEI NUMBER]": extracted_info.get("IMEI", ""),
                    "[MODEL]": extracted_info.get("MODEL", ""),
                    "[PHONE NUMBER]": phone_number
                }

                # Create a new document for each QR code, replace placeholders, and save
                modified_doc = replace_placeholders_in_docx(template_path, placeholders)
                output_doc_path = os.path.join(output_folder_path, f"output_{filename.replace('.png', '.docx')}")
                modified_doc.save(output_doc_path)
                print(f"Document saved: {output_doc_path}")
            else:
                print(f"QR code extraction failed or returned an error: {extracted_info}")
